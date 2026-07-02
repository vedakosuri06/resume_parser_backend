import json
import pandas as pd
import io
from docx import Document

def to_json(data: dict) -> str:
    return json.dumps(data, indent=2)


def to_excel(data: dict) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        basic = {k: [v] for k, v in data.items() if k not in ("skills", "experience", "education", "projects", "certifications", "achievements", "links")}
        pd.DataFrame(basic).to_excel(writer, sheet_name="Basic Info", index=False)
        pd.DataFrame({"Skills": data.get("skills", [])}).to_excel(writer, sheet_name="Skills", index=False)
        pd.DataFrame({"Experience": [data.get("experience", "")]}).to_excel(writer, sheet_name="Experience", index=False)
        pd.DataFrame({"Projects": [data.get("projects", "")]}).to_excel(writer, sheet_name="Projects", index=False)
        pd.DataFrame({"Certifications": [data.get("certifications", "")]}).to_excel(writer, sheet_name="Certifications", index=False)
        pd.DataFrame({"Achievements": [data.get("achievements", "")]}).to_excel(writer, sheet_name="Achievements", index=False)
        pd.DataFrame({"Education": [data.get("education", "")]}).to_excel(writer, sheet_name="Education", index=False)
        links = data.get("links", {})
        pd.DataFrame({"Platform": list(links.keys()), "URL": list(links.values())}).to_excel(writer, sheet_name="Links", index=False)
    return output.getvalue()


def to_word(data: dict) -> bytes:
    doc = Document()
    doc.add_heading("Resume - Parsed Data", 0)

    doc.add_heading("Basic Info", level=1)
    doc.add_paragraph(f"Name:  {data.get('name', 'N/A')}")
    doc.add_paragraph(f"Email: {data.get('email', 'N/A')}")
    doc.add_paragraph(f"Phone: {data.get('phone', 'N/A')}")

    links = data.get("links", {})
    if links:
        doc.add_heading("Links", level=1)
        for k, v in links.items():
            doc.add_paragraph(f"{k.capitalize()}: {v}")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(data.get("skills", [])) or "N/A")

    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph(data.get("experience", "None"))

    doc.add_heading("Projects", level=1)
    doc.add_paragraph(data.get("projects", "None"))

    doc.add_heading("Certifications", level=1)
    doc.add_paragraph(data.get("certifications", "None"))

    doc.add_heading("Achievements & Activities", level=1)
    doc.add_paragraph(data.get("achievements", "None"))

    doc.add_heading("Education", level=1)
    doc.add_paragraph(data.get("education", "N/A"))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def to_txt(data: dict) -> str:
    links = data.get("links", {})
    links_str = "\n".join(f"  {k.capitalize()}: {v}" for k, v in links.items()) or "  None"
    lines = [
        "===== RESUME PARSED DATA =====",
        "",
        "--- Basic Info ---",
        f"Name:  {data.get('name', 'N/A')}",
        f"Email: {data.get('email', 'N/A')}",
        f"Phone: {data.get('phone', 'N/A')}",
        "",
        "--- Links ---",
        links_str,
        "",
        "--- Skills ---",
        ", ".join(data.get("skills", [])) or "None",
        "",
        "--- Work Experience ---",
        data.get("experience", "None"),
        "",
        "--- Projects ---",
        data.get("projects", "None"),
        "",
        "--- Certifications ---",
        data.get("certifications", "None"),
        "",
        "--- Achievements & Activities ---",
        data.get("achievements", "None"),
        "",
        "--- Education ---",
        data.get("education", "N/A"),
    ]
    return "\n".join(lines)